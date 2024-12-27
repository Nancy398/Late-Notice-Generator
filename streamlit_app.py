import streamlit as st
from io import BytesIO
import pypandoc
from num2words import num2words  # 用于将数字转换为英文大写
from datetime import datetime
from docx import Document
import pandoc
import fitz  # PyMuPDF
import os

# 加载本地模板
import streamlit as st
import fitz  # PyMuPDF
import os

def fill_pdf(output_path, data,text_parts):
    pdf = fitz.open("Late Notice.pdf")

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        for key, value in data.items():
            value = str(value)  # 确保值是字符串
            search_term = f"{{{{{key}}}}}"  # 占位符格式

            # 查找占位符位置
            matches = page.search_for(search_term)
            for match in matches:
                rect = match  # 获取占位符的矩形区域
                # 用白色矩形覆盖占位符区域
                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                is_bold = False
                is_underlined = False

                for part in text_parts:
                    if part[0] == key:
                        is_bold = part[1]
                        is_underlined = part[2]
                fontname = "timesb" if is_bold else "times"

                # 插入新的文本，确保字体大小适应
                page.insert_text(
                    (rect.x0,rect.y1-2.5),  # 插入文本的位置是占位符的左上角
                    value,
                    fontsize=12,  # 自动计算的字体大小
                    fontname=fontname,  # 字体名称
                    color=(0, 0, 0)  # 黑色文本
                )
                if is_underlined:
                    text_width = fitz.get_text_length(value, fontsize=fontsize, fontname=fontname)
                    underline_y = rect.y1 -1。5  # 下划线稍微低于文本基线
                    page.draw_line(
                        (rect.x0, underline_y),  # 起点
                        (rect.x0 + text_width, underline_y),  # 终点
                        color=(0, 0, 0),  # 黑色线条
                        width=0.5  # 下划线宽度
                    )

    # 保存修改后的 PDF
    pdf.save(output_path)
    pdf.close()


# Streamlit 界面
st.title("PDF 模板填充器")
last_name = st.text_input("Last Name")
full_name = st.text_input("Full Name")
address = st.text_area("Address")
postal = st.text_input("Postal Code")
title = st.selectbox("Title", ["Mr.", "Ms."])
amount = st.number_input("Amount", min_value=0.0, format="%.2f")
formatted_amount = "{:,.2f}".format(amount)

import inflect

def format_amount_in_words(amount):
    """
    将金额格式化为英文形式，并转换为单词形式。 
    例如 1234.56 转换为 'One Thousand Two Hundred Thirty-Four Dollars and 56/100 Cents'
    
    :param amount: float, 输入的金额
    :return: str, 格式化后的金额字符串
    """
    p = inflect.engine()

    # 拆分整数部分和小数部分
    dollars = int(amount)
    cents = int(round((amount - dollars) * 100))

    # 将整数部分转化为英文单词
    dollar_words = p.number_to_words(dollars).replace(",", "")  # 移除逗号
    
    # 将金额转换为每个单词首字母大写
    dollar_words = dollar_words.title()  # 移除逗号
    
    # 创建最终格式化字符串
    if cents == 0:
        return f"{dollar_words} Dollars"
    else:
        return f"{dollar_words} Dollars and {cents}/100 Cents"
amount_word = format_amount_in_words(amount)        
def amount_to_words(amount):
    return num2words(amount, to='currency', lang='en', currency ='USD').title()
amount_words = amount_to_words(amount)

def get_current_date():
    now = datetime.now()
    return now.strftime("%B %d, %Y")
current_date = get_current_date()
    
data = {
    "Full Name": full_name,
    "Last Name": last_name,
    "Address": address,
    "Postal": postal,
    "gen": title,
    "Amount Words":f"{amount_word} (${str(formatted_amount)}).",
    "Date":current_date,
    "DateB":current_date,
    "Full Address":f"{address}, Los Angeles, CA, {postal}",
    "gender": f"{title}{last_name},"
}

text_parts = [
    ("Full Name", False, False),
    ("Last Name", False, False),
    ("Address", False, False),
    ("Postal", False, False),
    ("Amount Words", True, True),
    ("Date", False, False),
    ("DateB", True, False),
    ("Full Address", True, False),
    ("gender", False, False),
]

# 生成 PDF
if st.button("生成 PDF"):
        output_path = "filled_template.pdf"
        fill_pdf(output_path, data,text_parts)
        with open(output_path, "rb") as f:
            st.download_button(
                label="下载 PDF",
                data=f,
                file_name="filled_template.pdf",
                mime="application/pdf"
            )
        st.success("PDF 生成成功！")
#     else:
#         st.error("请先上传 PDF 模板！")


# # 用户输入替换值

# # 将金额转换为英文大写


# if st.button("生成 PDF"):
#     for paragraph in doc.paragraphs:
#         if "{First Name}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{First Name}", first_name)
#         if "{Last Name}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{Last Name}", last_name)
#         if "{Date}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{Date}", current_date)
#         if "{Address}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{Address}", address)
#         if "{Postal}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{Postal}", postal)
#         if "{gender}" in paragraph.text:
#             paragraph.text = paragraph.text.replace("{gender}", title)
#         if "{Amount}" in paragraph.text:
#             before_text = paragraph.text.split("{Amount}")[0]
#             after_text = paragraph.text.split("{Amount}")[1]
#             paragraph.clear()
#             paragraph.add_run(before_text)
#             run = paragraph.add_run(str(formatted_amount))
#             run.bold = True
#             run.underline = True
#             paragraph.add_run(after_text)
#         if "{Amount Words}" in paragraph.text:
#             before_text = paragraph.text.split("{Amount Words}")[0]
#             after_text = paragraph.text.split("{Amount Words}")[1]
#             paragraph.clear()
#             paragraph.add_run(before_text)
#             run = paragraph.add_run(amount_words)
#             run.bold = True
#             run.underline = True
#             paragraph.add_run(after_text)

# doc.save("Transfer.docx")
# #     # 提供下载链接
# st.success("Docx 已生成！")
# with open("Transfer.docx", "rb") as f:
#     st.download_button(
#         label="下载 docx 文件",
#         data=f,
#         file_name="Late Notice.docx",
#         mime="application/pdf"
#     )



# Title of the Streamlit app
# st.title("DOCX to PDF Converter")
# pypandoc.download_pandoc()
# # File uploader
# uploaded_file = st.file_uploader("Upload DOCX file", type="docx")

# if uploaded_file is not None:
#     # Save the uploaded DOCX file to BytesIO object
#     docx_file = BytesIO(uploaded_file.read())

#     # Convert the DOCX to PDF using pypandoc
#     try:
#         # We need to save the DOCX file to a temporary file first for pypandoc to work
#         temp_input_path = "/tmp/uploaded_file.docx"
#         with open(temp_input_path, 'wb') as f:
#             f.write(docx_file.getbuffer())
        
#         # Output PDF path
#         temp_output_path = "/tmp/output.pdf"
        
#         # Convert the DOCX file to PDF using pypandoc
#         pypandoc.convert_file(temp_input_path, to='pdf', format='docx', outputfile=temp_output_path,extra_args=['--pdf-engine=xelatex'])
        
#         # Provide download link for the converted PDF
#         with open(temp_output_path, "rb") as pdf_file:
#             st.download_button("Download PDF", pdf_file, file_name="converted_file.pdf", mime="application/pdf")
        
#     except Exception as e:
#         st.error(f"Error: {e}")



